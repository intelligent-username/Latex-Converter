# Used by Microsoft word 2007 and onwards and Google Docs

from docx import Document
from bs4 import BeautifulSoup, Tag

def convert_to_html(fname):
    output_file = fname.rsplit('.', 1)[0] + '.html'
    
    doc = Document(fname)
    soup = BeautifulSoup('<html><body></body></html>', 'html.parser')
    body = soup.body
    
    for para in doc.paragraphs:
        p = soup.new_tag('p')
        for run in para.runs:
            if run.bold:
                b = soup.new_tag('b')
                b.string = run.text
                p.append(b)
            elif run.italic:
                i = soup.new_tag('i')
                i.string = run.text
                p.append(i)
            else:
                p.append(run.text)
        body.append(p)
    
    for table in doc.tables:
        t = soup.new_tag('table')
        for row in table.rows:
            tr = soup.new_tag('tr')
            for cell in row.cells:
                td = soup.new_tag('td')
                td.string = cell.text
                tr.append(td)
            t.append(tr)
        body.append(t)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(str(soup))
    
    print(f"Converted {fname} to HTML successfully.")
    return output_file
